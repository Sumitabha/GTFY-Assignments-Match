import azure.functions as func
import logging
import os
import tempfile
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from openai import AzureOpenAI
import json
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

enhanceCV = func.Blueprint()

def extract_resume_text(filepath):
    ext = os.path.splitext(filepath)[-1].lower()
    if ext == ".pdf":
        import fitz
        with fitz.open(filepath) as doc:
            return "\n".join(page.get_text() for page in doc)
    elif ext == ".docx":
        from docx import Document
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".txt":
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError("Unsupported file type: " + ext)

def get_latest_resume_from_folder(blob_service_client, container_name, folder_prefix):
    container_client = blob_service_client.get_container_client(container_name)
    blobs = list(container_client.list_blobs(name_starts_with=folder_prefix))
    if not blobs:
        raise FileNotFoundError(f"No files found in folder: {folder_prefix}")
    latest_blob = max(blobs, key=lambda b: b.last_modified)
    return latest_blob.name

def save_text_to_docx(text, path):
    doc = Document()

    # Create and configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Remove spacing between paragraphs
    for para_format in [style.paragraph_format]:
        para_format.space_after = Pt(0)
        para_format.space_before = Pt(0)
        para_format.line_spacing = 1.0

    for line in text.splitlines():
        line = line.strip()

        if not line or line == "---":
            continue  # skip unnecessary blank lines or dividers

        # Headings detection
        if line.startswith("###"):
            doc.add_paragraph(line.replace("###", "").strip(), style='Heading 2')
        elif line.startswith("**") and line.endswith("**"):
            doc.add_paragraph(line.strip("*"), style='Heading 1')
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(path)

@enhanceCV.route(route="enhanceCV", methods=["POST"])
def enhanceResume(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Processing resume enhancement.')

    try:
        job = req.get_json()
        job_description = job.get("job_desc", "")
        required_skills = job.get("req_skills", "")

        if not job_description:
            return func.HttpResponse("Job description is required.", status_code=400)

        # Blob setup
        container_name = "gtfydemo"
        resume_folder = "resume"
        output_folder = "enhanced_cv"

        blob_service = BlobServiceClient.from_connection_string(os.environ["AZURE_BLOB_CONN"])
        blob_name = get_latest_resume_from_folder(blob_service, container_name, resume_folder)
        blob_client = blob_service.get_blob_client(container=container_name, blob=blob_name)

        # Download resume
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(blob_name)[-1]) as tmp_file:
            tmp_file.write(blob_client.download_blob().readall())
            tmp_resume_path = tmp_file.name

        resume_text = extract_resume_text(tmp_resume_path)
        os.remove(tmp_resume_path)

        # Call GPT-4o
        client = AzureOpenAI(
            api_key=os.environ["AZURE_OPENAI_KEY"],
            api_version=os.environ["AZURE_OPENAI_API_VERSION"],
            azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"]
        )

        prompt = f"""
        You are a resume expert. Rewrite the following resume to be better tailored for this job.

        Job Description:
        {job_description}

        Required Skills:
        {required_skills}

        Resume:
        {resume_text}
        """

        response = client.chat.completions.create(
            model=os.environ["AZURE_OPENAI_DEPLOYMENT"],
            messages=[
                {"role": "system", "content": "You are an AI resume optimization assistant."},
                {"role": "user", "content": prompt}
            ]
        )

        enhanced_resume_text = response.choices[0].message.content.strip()

        # Save as .docx
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as enhanced_doc:
            save_text_to_docx(enhanced_resume_text, enhanced_doc.name)
            enhanced_doc_path = enhanced_doc.name

        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        output_blob_name = f"{output_folder}/enhanced_resume_{timestamp}.docx"
        output_blob_client = blob_service.get_blob_client(container=container_name, blob=output_blob_name)

        with open(enhanced_doc_path, "rb") as f:
            output_blob_client.upload_blob(f, overwrite=True)

        os.remove(enhanced_doc_path)

        # Generate SAS URL
        sas_url = generate_blob_sas(
            account_name=blob_service.account_name,
            container_name=container_name,
            blob_name=output_blob_name,
            account_key=blob_service.credential.account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(minutes=60)
        )

        download_url = f"https://{blob_service.account_name}.blob.core.windows.net/{container_name}/{output_blob_name}?{sas_url}"

        return func.HttpResponse(
            json.dumps({
                "message": "Enhanced resume uploaded successfully.",
                "download_url": download_url
            }),
            mimetype="application/json",
            status_code=200
        )

    except Exception as e:
        logging.exception("Error during resume enhancement.")
        return func.HttpResponse(f"Error: {str(e)}", status_code=500)
